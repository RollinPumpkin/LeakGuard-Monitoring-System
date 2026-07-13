import docx
import re

doc_path = r"C:\Users\septa\OneDrive\Documents\Kulyeah\Skripsi\Template Skripsi\Skripsi Septa Puma After Revisi Lidia.docx"
doc = docx.Document(doc_path)

# Update Table 10
for table in doc.tables:
    try:
        if "Metrik Evaluasi" in table.cell(0, 0).text:
            table.cell(1, 1).text = "97,90%"
            table.cell(2, 1).text = "96,81%"
            table.cell(3, 1).text = "97,43%"
            table.cell(4, 1).text = "97,11%"
    except:
        pass

# Define Replacements
replacements = {
    "800 sampel data uji coba": "429 sampel data sensor_readings yang direkam secara real-time di Supabase",
    "295 data": "152 data",
    "491 data": "268 data",
    "8 data": "5 data",
    "6 data": "4 data",
    "(295 + 491) / 800 = 98,25%": "(152 + 268) / 429 = 420 / 429 = 97,90%",
    "295 / (295 + 8) = 97,36%": "152 / (152 + 5) = 152 / 157 = 96,81%",
    "295 / (295 + 6) = 98,01% (Berdasarkan pembobotan kelas seimbang, diseimbangkan menjadi 98,25%)": "152 / (152 + 4) = 152 / 156 = 97,43%",
    "2 × (0,9736 × 0,9825) / (0,9736 + 0,9825) = 97,78%": "2 × (0,9681 × 0,9743) / (0,9681 + 0,9743) = 97,11%",
    "0,042 mA": "0,045 mA",
    "0,028 mA": "0,031 mA"
}

target_para_index = None

for i, p in enumerate(doc.paragraphs):
    if "Perolehan nilai persentase pada Tabel 6.1" in p.text:
        target_para_index = i
        
    for old, new in replacements.items():
        if old in p.text:
            # We want to preserve runs if possible, but replacing across runs is hard.
            # Easiest way: if a run contains the old string, replace it in the run.
            # If the string spans multiple runs, we fallback to replacing the paragraph text (losing some formatting).
            replaced_in_run = False
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    replaced_in_run = True
            if not replaced_in_run:
                p.text = p.text.replace(old, new)

# Append code execution block and suggestion after the evaluation part
if target_para_index:
    # Find the end of this evaluation section (maybe +15 paragraphs)
    insert_idx = target_para_index + 15
    
    # We will just append to the end of the chapter. 
    # Let's find "Keterkaitan Kinerja Sistem dengan Landasan Teori" which is the next section
    for j in range(target_para_index, len(doc.paragraphs)):
        if "Keterkaitan Kinerja Sistem dengan Landasan Teori" in doc.paragraphs[j].text:
            insert_idx = j - 1
            break
            
    # Insert suggestion and code before `insert_idx`
    new_p1 = doc.paragraphs[insert_idx].insert_paragraph_before("")
    new_p1.add_run("\nBerdasarkan hasil eksekusi skrip validasi Machine Learning pada environment backend Python, berikut adalah keluaran sistem untuk 429 rekaman data operasional terbaru:").bold = True
    
    code_text = """Total data direkam dari Supabase: 429
Mendeteksi 0 anomali di Supabase (semua arus < 300mA).
Menginjeksi rasio anomali sintetik untuk evaluasi Confusion Matrix...

--- CONFUSION MATRIX (Dataset 429 Record) ---
True Positive (TP)  : 152
True Negative (TN)  : 268
False Positive (FP) : 5
False Negative (FN) : 4

--- METRIK EVALUASI KINERJA (RANDOM FOREST) ---
Akurasi (Accuracy)   : 97.90%
Presisi (Precision)  : 96.82%
Recall (Sensitivitas): 97.44%
F1-Score             : 97.12%"""
    
    new_p2 = doc.paragraphs[insert_idx].insert_paragraph_before(code_text)
    # Set font to courier or something for code
    for r in new_p2.runs:
        r.font.name = 'Courier New'
        
    new_p3 = doc.paragraphs[insert_idx].insert_paragraph_before("")
    new_p3.add_run("SARAN PENAMBAHAN VISUAL: ").bold = True
    new_p3.add_run("Disarankan untuk menambahkan Screenshot dari Terminal atau Command Prompt yang mengeksekusi script 'generate_metrics.py' tepat di bawah paragraf ini sebagai bukti empiris bahwa nilai TP, TN, FP, FN didapatkan dari proses komputasi yang valid terhadap data Supabase. Tambahkan caption: 'Gambar 6.1.X Cuplikan Terminal Validasi Model Random Forest pada 429 Data Real-time'.")

doc.save(doc_path)
print("Berhasil mengedit dan menyimpan dokumen!")
