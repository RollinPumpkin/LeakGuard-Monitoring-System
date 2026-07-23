import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc_path = r"C:\Users\septa\OneDrive\Documents\Kulyeah\Skripsi\Template Skripsi\Skripsi Septa Puma ENDGAME.docx"
doc = docx.Document(doc_path)

# Data for tables
tables_data = [
    [
        ["Judul", "Machine Learning Pipeline for Battery State-of-Health Estimation"],
        ["Masalah", "Kurangnya metode prediksi kesehatan baterai (State-of-Health) yang akurat dan tahan terhadap data pencilan (outlier) pada kondisi operasional dinamis."],
        ["Desain eksperimen", "Peneliti membangun pipeline Machine Learning berbasis model deret waktu (time-series forecasting) untuk memonitor siklus hidup baterai. Data operasional diekstraksi dan diproses menggunakan algoritma prediktif mutakhir."],
        ["Hasil", "Pipeline tersebut berhasil meningkatkan akurasi estimasi sisa umur baterai secara drastis, membuktikan bahwa algoritma forecasting (yang sejenis dengan penelitian ini) sangat tangguh dan bisa diandalkan untuk memprediksi keandalan alat sebelum terjadi kegagalan sistem."]
    ],
    [
        ["Judul", "Insulator leakage current prediction based on generative adversarial networks and optimized support vector regression with crisscross optimization algorithm"],
        ["Masalah", "Sulitnya memprediksi arus bocor secara akurat karena data lapangan yang seringkali tidak seimbang dan dipenuhi oleh noise (gangguan) dari lingkungan sekitar."],
        ["Desain eksperimen", "Peneliti menggunakan metode pemantauan arus bocor isolator secara langsung, namun memproses datanya menggunakan pendekatan gabungan Support Vector Regression (SVR) dan Generative Adversarial Networks (GAN)."],
        ["Hasil", "Algoritma SVR yang dioptimasi terbukti mampu melakukan pemodelan regresi yang akurat untuk arus bocor. Ini membuktikan bahwa metode pemantauan (monitoring) arus bocor yang sama dapat diselesaikan dengan kelas algoritma komputasi lain (SVR/GAN), sehingga memberikan komparasi dan memperkaya justifikasi pemilihan metodologi pada penelitian ini."]
    ]
]

bib_data = [
    "Roman, D., Saxena, S., Robu, V., Pecht, M., & Flynn, D. (2021). Machine learning pipeline for battery state-of-health estimation. Nature Machine Intelligence, 3(5), 447-456. https://doi.org/10.1038/s42256-021-00312-3",
    "Wen, H., Zhang, J., Wen, H., Wu, J., Zhao, X., Lin, W., & Zhang, H. (2022). Insulator leakage current prediction based on generative adversarial networks and optimized support vector regression with crisscross optimization algorithm. 2022 IEEE International Conference on Electrical Engineering and Computer Technologies (CEECT). https://doi.org/10.1109/ceect55960.2022.10030195"
]

target_para = None
bib_para = None

for para in doc.paragraphs:
    if "Penambahan Related Research (Algoritma sama metode beda" in para.text:
        target_para = para
    if "DAFTAR PUSTAKA" in para.text.upper():
        bib_para = para

def style_run(run):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

if target_para:
    prev_element = target_para._p
    
    # Text introduction for first paper
    p1 = doc.add_paragraph()
    run = p1.add_run("Tabel berikut menyajikan penelitian dengan algoritma yang sama namun metode yang berbeda:")
    style_run(run)
    prev_element.addnext(p1._p)
    prev_element = p1._p

    # Table 1
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = 'Table Grid'
    for i, row_data in enumerate(tables_data[0]):
        row = t1.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            style_run(r)
    prev_element.addnext(t1._tbl)
    prev_element = t1._tbl

    # Text introduction for second paper
    p2 = doc.add_paragraph()
    run2 = p2.add_run("\nTabel berikut menyajikan penelitian dengan metode yang sama namun algoritma yang berbeda:")
    style_run(run2)
    prev_element.addnext(p2._p)
    prev_element = p2._p

    # Table 2
    t2 = doc.add_table(rows=4, cols=2)
    t2.style = 'Table Grid'
    for i, row_data in enumerate(tables_data[1]):
        row = t2.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            style_run(r)
    prev_element.addnext(t2._tbl)
    prev_element = t2._tbl

# Add to Bibliography
# If we couldn't find DAFTAR PUSTAKA header, we just append at the end
for bib in bib_data:
    p_bib = doc.add_paragraph()
    run_bib = p_bib.add_run(bib)
    style_run(run_bib)
    
    # We want hanging indent for APA style
    p_format = p_bib.paragraph_format
    p_format.left_indent = Pt(36)
    p_format.first_line_indent = Pt(-36)

doc.save(doc_path)
print("Docx updated successfully!")
