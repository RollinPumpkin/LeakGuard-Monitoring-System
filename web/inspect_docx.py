import docx
import sys

doc = docx.Document(r"C:\Users\septa\OneDrive\Documents\Kulyeah\Skripsi\Template Skripsi\Skripsi Septa Puma After Revisi Lidia.docx")

print("--- TABLES ---")
for i, table in enumerate(doc.tables):
    try:
        first_cell = table.cell(0, 0).text.strip()
        print(f"Table {i}: First Cell = {first_cell[:30]}")
    except:
        pass

print("\n--- PARAGRAPHS WITH '800' atau 'Tabel 6.1' ---")
for i, p in enumerate(doc.paragraphs):
    if '800' in p.text or 'Tabel 6.1' in p.text or 'Random Forest' in p.text:
        print(f"Para {i}: {p.text[:100]}")
