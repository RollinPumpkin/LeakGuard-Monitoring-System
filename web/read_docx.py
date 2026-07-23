import docx
import json

doc_path = r"C:\Users\septa\OneDrive\Documents\Kulyeah\Skripsi\Template Skripsi\Skripsi Septa Puma ENDGAME.docx"

doc = docx.Document(doc_path)
content = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        content.append(f"[{i}] {para.text}")

# Let's save the first 500 paragraphs to a text file to avoid huge console output, or we can just print paragraphs that contain "Bab 3", "Related", "Penambahan", etc.
for c in content:
    if "Bab 3" in c or "BAB 3" in c or "BAB III" in c or "Penambahan" in c or "Related" in c:
        print(c)

# Let's also check tables
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}:")
    for j, row in enumerate(table.rows):
        row_text = [cell.text.strip() for cell in row.cells]
        print(f"Row {j}: {row_text}")
        if j > 2: # just print first few rows of each table
            break
