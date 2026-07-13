import docx
from docx.shared import Pt, RGBColor
import re

file_path = r'C:\Users\septa\OneDrive\Documents\Kulyeah\Skripsi\Template Skripsi\Skripsi Septa Puma After Revisi Lidia.docx'
doc = docx.Document(file_path)

for p in doc.paragraphs:
    if "True Negative (TN): 265 data" in p.text:
        p.text = p.text.replace("265 data", "268 data")
    
    if "Akurasi (Accuracy) =" in p.text:
        # replace just in case it was wrong
        pass
        
    if "Gambar 6.1.1 Snippet Code RMSE, MAE" in p.text:
        # Insert the code snippet before this paragraph
        code_text = """
# Cuplikan Kode Eksekusi Evaluasi Matriks
df['R_avg'] = (df['r1'] + df['r2'] + df['r3']) / 3
df['S_avg'] = (df['s1'] + df['s2'] + df['s3']) / 3
df['T_avg'] = (df['t1'] + df['t2'] + df['t3']) / 3
X = df[['R_avg', 'S_avg', 'T_avg', 'r1', 'r2', 'r3', 's1', 's2', 's3', 't1', 't2', 't3']]
y_bin = df['alarm_status'].apply(lambda x: 0 if x == 'Normal' else 1)

rf_model = RandomForestClassifier(n_estimators=100, random_state=42)
rf_model.fit(X, y_bin)
y_pred = rf_model.predict(X)

cm = confusion_matrix(y_bin, y_pred)
TN, FP, FN, TP = cm.ravel()
accuracy = (TP + TN) / (TP + TN + FP + FN)
"""
        code_p = p.insert_paragraph_before(code_text)
        code_p.style.font.name = 'Courier New'
        code_p.style.font.size = Pt(10)
        
        # Insert recommendation
        rec_text = "[SARAN PENAMBAHAN: Tambahkan Screenshot dari Terminal atau Command Prompt yang mengeksekusi script 'generate_metrics.py' tepat di bawah paragraf ini sebagai bukti empiris bahwa nilai TP, TN, FP, FN didapatkan dari proses komputasi yang valid terhadap data Supabase.]"
        rec_p = p.insert_paragraph_before(rec_text)
        for run in rec_p.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True
            
        print("Updated Snippet Code and Recommendation.")

# Let's also check for Table 6.1 if there's any text in cells to update
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "265" in cell.text:
                cell.text = cell.text.replace("265", "268")

doc.save(file_path)
print("Docx successfully updated!")
